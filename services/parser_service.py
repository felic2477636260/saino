import hashlib
import logging
import re
from dataclasses import dataclass
from pathlib import Path

import fitz


logger = logging.getLogger(__name__)


@dataclass
class ParsedPage:
    page_no: int
    text: str


@dataclass
class ParsedDocument:
    doc_id: str
    company_code: str
    company_name: str
    industry: str
    report_type: str
    filename: str
    source_path: str
    total_pages: int
    title: str
    pages: list[ParsedPage]


class PDFParserService:
    COMPANY_MAP = {
        "吉比特": ("603444", "game"),
        "三七互娱": ("002555", "game"),
        "完美世界": ("002624", "game"),
    }
    SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx"}

    def parse_file(self, file_path: Path) -> ParsedDocument:
        suffix = file_path.suffix.lower()
        if suffix == ".pdf":
            return self.parse_pdf(file_path)
        if suffix in {".txt", ".md"}:
            return self._parse_text_document(file_path, report_type=self._infer_report_type(file_path))
        if suffix == ".docx":
            return self._parse_docx_document(file_path, report_type=self._infer_report_type(file_path))
        raise ValueError(f"unsupported file type: {suffix}")

    def parse_pdf(self, file_path: Path) -> ParsedDocument:
        doc_id = hashlib.md5(str(file_path).encode("utf-8")).hexdigest()
        company_code, company_name, industry = self._infer_company(file_path)
        report_type = self._infer_report_type(file_path)
        pages: list[ParsedPage] = []
        title = file_path.stem

        pdf = fitz.open(file_path)
        try:
            for index, page in enumerate(pdf, start=1):
                try:
                    text = page.get_text("text").strip()
                except Exception as exc:
                    logger.exception("failed to extract page %s of %s: %s", index, file_path, exc)
                    text = ""
                pages.append(ParsedPage(page_no=index, text=text))
            metadata_title = (pdf.metadata or {}).get("title") or ""
            if metadata_title.strip():
                title = metadata_title.strip()
        finally:
            pdf.close()

        return ParsedDocument(
            doc_id=doc_id,
            company_code=company_code,
            company_name=company_name,
            industry=industry,
            report_type=report_type,
            filename=file_path.name,
            source_path=str(file_path),
            total_pages=len(pages),
            title=title,
            pages=pages,
        )

    def _parse_text_document(self, file_path: Path, *, report_type: str) -> ParsedDocument:
        text = self._read_text_file(file_path)
        return self._build_single_page_document(file_path=file_path, text=text, report_type=report_type)

    def _parse_docx_document(self, file_path: Path, *, report_type: str) -> ParsedDocument:
        from docx import Document

        document = Document(file_path)
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        text = "\n".join(paragraphs)
        return self._build_single_page_document(file_path=file_path, text=text, report_type=report_type)

    def _build_single_page_document(self, *, file_path: Path, text: str, report_type: str) -> ParsedDocument:
        doc_id = hashlib.md5(str(file_path).encode("utf-8")).hexdigest()
        company_code, company_name, industry = self._infer_company(file_path)
        title = file_path.stem
        pages = [ParsedPage(page_no=1, text=text.strip())]
        return ParsedDocument(
            doc_id=doc_id,
            company_code=company_code,
            company_name=company_name,
            industry=industry,
            report_type=report_type,
            filename=file_path.name,
            source_path=str(file_path),
            total_pages=1,
            title=title,
            pages=pages,
        )

    def _infer_company(self, file_path: Path) -> tuple[str, str, str]:
        path_text = str(file_path)
        for name, (code, industry) in self.COMPANY_MAP.items():
            if name in path_text:
                return code, name, industry

        folder_match = re.search(r"(?:company|research|industry|custom)[\\/](\w+)", path_text, re.IGNORECASE)
        if folder_match:
            code = folder_match.group(1)
            return code, code, "generic"
        return "macro", "行业材料", "generic"

    def _infer_report_type(self, file_path: Path) -> str:
        path_text = str(file_path).lower()
        name = file_path.name.lower()
        if "research" in path_text or "研报" in file_path.name or "研究" in file_path.name:
            return "research_report"
        if "industry" in path_text or "行业" in file_path.name:
            return "industry_report"
        if "年报" in file_path.name or "annual" in name or "季报" in file_path.name or "quarter" in name:
            return "annual_report"
        return "annual_report"

    @staticmethod
    def _read_text_file(file_path: Path) -> str:
        encodings = ("utf-8", "utf-8-sig", "gbk", "gb18030")
        for encoding in encodings:
            try:
                return file_path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue
        return file_path.read_text(encoding="utf-8", errors="ignore")
